# -*- coding: utf-8 -*-
"""Izvoz rasporeda: Excel (po nastavniku + po razredu) i narativni izvještaj (.docx)."""
import collections
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    HAS_DOCX=True
except Exception:
    HAS_DOCX=False

PAL={"navy":"035EA1","blue":"08ABE6","yellow":"FFCB29","red":"E8262C","black":"000000"}

def _grid(lessons, key):
    g=collections.defaultdict(dict)
    for L in lessons: g[L[key]][(L["day"],L["period"])]=L
    return g

def _grid_multi(lessons, key):
    """Kao _grid, ali SVAKA ćelija čuva LISTU časova — potrebno za pogled
    po nastavniku kad isti termin dijele dva razreda (npr. spojeni Sport
    7&8): bez ovoga se drugi razred tiho izgubi (posljednji u listi
    prepisuje prethodni)."""
    g=collections.defaultdict(lambda: collections.defaultdict(list))
    for L in lessons: g[L[key]][(L["day"],L["period"])].append(L)
    return g

def _cell_label(cell_lessons):
    """Spoji jednu ili više lekcija u istom terminu u jedan tekst.
    Isti predmet, različiti razredi (spojeni čas) -> 'Sport 7/8'.
    Različit predmet u istom terminu bi bio grešaka koju validator hvata
    prije ovoga, ali ovdje se svejedno prikazuje sve, ništa se ne gubi."""
    if not cell_lessons: return ""
    by_subj = collections.defaultdict(list)
    for L in cell_lessons:
        if L["grade"] is not None: by_subj[L["subj"]].append(L["grade"])
    parts=[]
    for subj, grades in by_subj.items():
        gtxt = "/".join(str(g) for g in sorted(set(grades)))
        parts.append(f"{subj} {gtxt}" if gtxt else subj)
    # predmeti bez razreda (npr. Schulbibliothek)
    for L in cell_lessons:
        if L["grade"] is None: parts.append(L["subj"])
    return " + ".join(parts)

def export_excel(lessons, config, path):
    D=config["days"]; P=config["periods"]; T=config.get("period_times",{})
    HR={int(k):v for k,v in config["homeroom"].items()}
    thin=Side(style="thin",color="B0B0B0"); B=Border(thin,thin,thin,thin)
    HF=PatternFill("solid",fgColor=PAL["navy"]); DF=PatternFill("solid",fgColor=PAL["blue"])
    NF=PatternFill("solid",fgColor="E2EFDA"); C=Alignment("center","center",wrap_text=True)
    wf=Font(name="Arial",size=10,bold=True,color="FFFFFF"); f=Font(name="Arial",size=10)
    wb=Workbook()
    # sheet 1: po nastavniku
    ws=wb.active; ws.title="NASTAVNICI"; ws.column_dimensions["A"].width=15
    for col in "BCDEF": ws.column_dimensions[col].width=19
    teachers=sorted(set(L["teacher"] for L in lessons))
    r=1; byt=_grid_multi(lessons,"teacher")
    for t in teachers:
        ws.merge_cells(start_row=r,start_column=1,end_row=r,end_column=1+len(D))
        cc=ws.cell(r,1,f"{t}   ·   {sum(1 for L in lessons if L['teacher']==t)} časova")
        cc.font=Font(name="Arial",bold=True,color="FFFFFF"); cc.fill=HF; cc.alignment=Alignment("left","center")
        for k in range(1,2+len(D)): ws.cell(r,k).fill=HF
        r+=1
        ws.cell(r,1,"Čas").font=wf; ws.cell(r,1).fill=DF; ws.cell(r,1).border=B; ws.cell(r,1).alignment=C
        for i,d in enumerate(D):
            x=ws.cell(r,2+i,d); x.font=wf; x.fill=DF; x.border=B; x.alignment=C
        r+=1
        for p in P:
            ws.cell(r,1,f"{p}. {T.get(str(p),'')}").font=f; ws.cell(r,1).border=B; ws.cell(r,1).alignment=C
            for i,d in enumerate(D):
                cell = byt[t].get((d,p), [])
                txt = _cell_label(cell)
                x=ws.cell(r,2+i,txt); x.font=f; x.border=B; x.alignment=C
                if any(L["subj"]=="Nacharbeit" for L in cell): x.fill=NF
            r+=1
        r+=1
    # sheet 2: po razredu
    ws2=wb.create_sheet("PO RAZREDIMA"); ws2.column_dimensions["A"].width=15
    for col in "BCDEF": ws2.column_dimensions[col].width=21
    r=1; byg=_grid(lessons,"grade")
    for g in config["grades"]:
        wk = sum(1 for L in lessons if L["grade"]==g)
        ws2.merge_cells(start_row=r,start_column=1,end_row=r,end_column=1+len(D))
        cc=ws2.cell(r,1,f"{g}. RAZRED   ·   razrednik: {HR.get(g,'')}   ·   {wk} časova sedmično")
        cc.font=Font(name="Arial",bold=True,color="FFFFFF"); cc.fill=HF; cc.alignment=Alignment("left","center")
        for k in range(1,2+len(D)): ws2.cell(r,k).fill=HF
        r+=1
        ws2.cell(r,1,"Čas").font=wf; ws2.cell(r,1).fill=DF; ws2.cell(r,1).border=B; ws2.cell(r,1).alignment=C
        for i,d in enumerate(D):
            x=ws2.cell(r,2+i,d); x.font=wf; x.fill=DF; x.border=B; x.alignment=C
        r+=1
        for p in P:
            ws2.cell(r,1,f"{p}. {T.get(str(p),'')}").font=f; ws2.cell(r,1).border=B; ws2.cell(r,1).alignment=C
            for i,d in enumerate(D):
                L=byg[g].get((d,p)); txt=""
                if L: txt=f"{L['subj']}\n{L['teacher'].split()[0]}"
                x=ws2.cell(r,2+i,txt); x.font=f; x.border=B; x.alignment=C
                if L and L["subj"]=="Nacharbeit": x.fill=NF
            ws2.row_dimensions[r].height=30
            r+=1
        r+=1
    wb.save(path); return path

def export_report(lessons, config, path):
    """Narativni izvještaj (.docx). Zahtijeva python-docx."""
    if not HAS_DOCX: raise RuntimeError("python-docx nije instaliran")
    HR={int(k):v for k,v in config["homeroom"].items()}
    doc=Document()
    h=doc.add_heading("IDSS — Raspored časova: narativni izvještaj",0)
    doc.add_paragraph("Automatski generisan raspored. Pravila: matematika i njemački u 1.–4. času; "
        "Nacharbeit uvijek na kraju dana (6./7.); svaki predmet najviše jednom dnevno po razredu, "
        "osim Sporta i likovnog (1.–4.) kao blok; razrednici imaju kontakt sa svojim razredom svaki dan.")
    doc.add_heading("Razredništva", level=1)
    for g in config["grades"]:
        wk = sum(1 for L in lessons if L["grade"]==g)
        doc.add_paragraph(f"{g}. razred — {HR.get(g,'')}   ·   {wk} časova sedmično", style="List Bullet")
    doc.add_heading("Zaduženja po nastavnicima", level=1)
    for t in sorted(set(L["teacher"] for L in lessons)):
        subs=collections.defaultdict(set)
        for L in lessons:
            if L["teacher"]==t and L["grade"] is not None: subs[L["subj"]].add(L["grade"])
        parts="; ".join(f"{s} ({','.join(map(str,sorted(gs)))})" for s,gs in sorted(subs.items()))
        p=doc.add_paragraph(); p.add_run(t).bold=True
        p.add_run(f" — {sum(1 for L in lessons if L['teacher']==t)} časova: {parts}")
    doc.save(path); return path
